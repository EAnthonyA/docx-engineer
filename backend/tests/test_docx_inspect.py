import io
import json
import tempfile
from pathlib import Path

import pytest
from docx import Document

from app.docx_inspect import compute_diff, summarize


def _make_docx(paragraphs: list[str]) -> str:
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(tmp.name)
    return tmp.name


def test_summarize_returns_valid_json():
    path = _make_docx(["Hello world", "Second paragraph", "Third paragraph"])
    result = summarize(path)
    data = json.loads(result)
    assert data["total_paragraphs"] >= 3
    assert data["non_empty_paragraphs"] >= 3


def test_compute_diff_unchanged():
    path = _make_docx(["Same text"])
    diff = compute_diff(path, path)
    assert diff["changed"] == 0
    assert all(e["status"] == "unchanged" for e in diff["entries"])


def test_compute_diff_detects_change():
    orig = _make_docx(["Original text", "Unchanged paragraph"])

    doc = Document(orig)
    for para in doc.paragraphs:
        if para.text == "Original text":
            para.clear()
            run = para.add_run("Modified text")
            run.bold = True
    import tempfile
    modified = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
    doc.save(modified.name)

    diff = compute_diff(orig, modified.name)
    assert diff["changed"] > 0
    changed = [e for e in diff["entries"] if e["status"] == "changed"]
    assert len(changed) >= 1


def test_compute_diff_detects_added_paragraphs():
    orig = _make_docx(["one", "two"])
    mod = _make_docx(["one", "two", "three"])
    diff = compute_diff(orig, mod)
    added = [e for e in diff["entries"] if e["status"] == "added"]
    assert len(added) == 1
    assert added[0]["before"] is None
    assert added[0]["after"]["text"] == "three"
    assert diff["changed"] == 1


def test_compute_diff_detects_removed_paragraphs():
    orig = _make_docx(["one", "two", "three"])
    mod = _make_docx(["one", "three"])
    diff = compute_diff(orig, mod)
    removed = [e for e in diff["entries"] if e["status"] == "removed"]
    assert len(removed) == 1
    assert removed[0]["after"] is None
    assert removed[0]["before"]["text"] == "two"


def test_compute_diff_replace_length_mismatch():
    # 1 original line replaced by 2 new lines: one paired "changed", one tail "added".
    orig = _make_docx(["alpha", "shared"])
    mod = _make_docx(["beta", "gamma", "shared"])
    diff = compute_diff(orig, mod)
    statuses = [e["status"] for e in diff["entries"]]
    assert statuses.count("changed") == 1
    assert statuses.count("added") == 1
    assert statuses.count("unchanged") == 1  # "shared"
    assert diff["total"] == 3
    assert diff["changed"] == 2
