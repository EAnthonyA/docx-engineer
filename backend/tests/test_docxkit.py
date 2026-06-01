"""Characterization tests for the sandbox DocxTools toolkit.

These pin the CURRENT behavior of docxkit so the upcoming readability refactor
can be proven to change nothing. Built before the refactor; must stay green.
"""

from docx import Document
from docx.shared import Pt, RGBColor

from docxkit import DocxTools


# ---------------------------------------------------------------------------
# helpers
# ---------------------------------------------------------------------------

def _para_with_runs(runs):
    """Build a doc whose single body paragraph is composed of explicit runs.

    runs: list of (text, fmt_dict) where fmt_dict may set bold/italic/underline.
    Returns (doc, paragraph).
    """
    doc = Document()
    p = doc.add_paragraph()
    for text, fmt in runs:
        r = p.add_run(text)
        for k, v in fmt.items():
            setattr(r, k, v)
    return doc, p


def _runs_snapshot(para):
    """List of (text, bool(bold)) for the paragraph's current runs."""
    return [(r.text, bool(r.bold)) for r in para.runs]


# ---------------------------------------------------------------------------
# format_tagged
# ---------------------------------------------------------------------------

def test_format_tagged_single_run():
    doc, p = _para_with_runs([("Hello <B>World<D> test", {})])
    count = DocxTools().format_tagged(doc, r"<B>(.*?)<D>", bold=True)
    assert count == 1
    assert p.text == "Hello World test"
    bold_runs = [r.text for r in p.runs if r.bold]
    assert bold_runs == ["World"]


def test_format_tagged_spans_run_boundary():
    doc, p = _para_with_runs([("Hel", {}), ("lo <B>Wor", {}), ("ld<D> test", {})])
    count = DocxTools().format_tagged(doc, r"<B>(.*?)<D>", bold=True)
    assert count == 1
    assert p.text == "Hello World test"
    assert [r.text for r in p.runs if r.bold] == ["World"]


def test_format_tagged_preserves_preexisting_bold():
    # Regression: a non-tagged word that was bold in the original must stay bold
    # after the tagged region is reformatted (the "Nuzudyta" bug).
    doc, p = _para_with_runs([
        ("<B>Tag<D> ", {}),
        ("Nuzudyta", {"bold": True}),
        (" end", {}),
    ])
    count = DocxTools().format_tagged(doc, r"<B>(.*?)<D>", bold=True)
    assert count == 1
    assert p.text == "Tag Nuzudyta end"
    snap = _runs_snapshot(p)
    # Tag is now bold; Nuzudyta keeps its original bold; surrounding stays plain.
    assert ("Nuzudyta", True) in snap
    assert any(text == "Tag" and bold for text, bold in snap)
    assert any("end" in text and not bold for text, bold in snap)


def test_format_tagged_no_match_leaves_paragraph_untouched():
    doc, p = _para_with_runs([("plain text here", {})])
    count = DocxTools().format_tagged(doc, r"<B>(.*?)<D>", bold=True)
    assert count == 0
    assert p.text == "plain text here"


def test_format_tagged_applies_color_and_size():
    doc, p = _para_with_runs([("a <B>word<D> b", {})])
    count = DocxTools().format_tagged(doc, r"<B>(.*?)<D>", color="FF0000", size_pt=20)
    assert count == 1
    word = [r for r in p.runs if r.text == "word"][0]
    assert word.font.color.rgb == RGBColor(0xFF, 0x00, 0x00)
    assert word.font.size == Pt(20)


def test_format_tagged_processes_tables_and_headers():
    doc = Document()
    table = doc.add_table(rows=1, cols=1)
    table.cell(0, 0).paragraphs[0].add_run("x <B>cell<D> y")
    doc.sections[0].header.paragraphs[0].add_run("h <B>head<D> z")

    count = DocxTools().format_tagged(doc, r"<B>(.*?)<D>", bold=True, locations="all")
    assert count == 2
    assert table.cell(0, 0).paragraphs[0].text == "x cell y"
    assert doc.sections[0].header.paragraphs[0].text == "h head z"


# ---------------------------------------------------------------------------
# replace_text
# ---------------------------------------------------------------------------

def test_replace_text_spans_run_boundary():
    doc, p = _para_with_runs([("foo", {}), ("bar", {})])
    count = DocxTools().replace_text(doc, "oob", "XYZ")
    assert count == 1
    assert p.text == "fXYZar"


def test_replace_text_locations_body_skips_header():
    doc = Document()
    doc.add_paragraph("hit here")
    doc.sections[0].header.paragraphs[0].add_run("hit here")

    count = DocxTools().replace_text(doc, "hit", "X", locations="body")
    assert count == 1
    assert doc.paragraphs[0].text == "X here"
    assert doc.sections[0].header.paragraphs[0].text == "hit here"


# ---------------------------------------------------------------------------
# regex_replace
# ---------------------------------------------------------------------------

def test_regex_replace_with_backrefs():
    doc, p = _para_with_runs([("01.02.2003", {})])
    count = DocxTools().regex_replace(doc, r"(\d{2})\.(\d{2})\.(\d{4})", r"\3-\2-\1")
    assert count == 1
    assert p.text == "2003-02-01"


def test_regex_replace_with_callable():
    doc, p = _para_with_runs([("abc def", {})])
    count = DocxTools().regex_replace(doc, r"abc", lambda m: m.group(0).upper())
    assert count == 1
    assert p.text == "ABC def"


# ---------------------------------------------------------------------------
# iter_paragraphs
# ---------------------------------------------------------------------------

def test_iter_paragraphs_covers_body_table_header():
    doc = Document()
    doc.add_paragraph("body")
    doc.add_table(rows=1, cols=1).cell(0, 0).paragraphs[0].add_run("cell")
    doc.sections[0].header.paragraphs[0].add_run("head")

    all_texts = [p.text for p in DocxTools().iter_paragraphs(doc, locations="all")]
    assert "body" in all_texts
    assert "cell" in all_texts
    assert "head" in all_texts

    body_texts = [p.text for p in DocxTools().iter_paragraphs(doc, locations="body")]
    assert "head" not in body_texts


# ---------------------------------------------------------------------------
# delete_paragraphs
# ---------------------------------------------------------------------------

def test_delete_paragraphs_removes_matching():
    doc = Document()
    doc.add_paragraph("keep")
    doc.add_paragraph("TODO: drop me")
    doc.add_paragraph("also keep")

    count = DocxTools().delete_paragraphs(doc, lambda t: t.startswith("TODO:"))
    assert count == 1
    assert [p.text for p in doc.paragraphs] == ["keep", "also keep"]


# ---------------------------------------------------------------------------
# set_format
# ---------------------------------------------------------------------------

def test_set_format_applies_to_all_runs():
    doc, p = _para_with_runs([("one ", {}), ("two", {})])
    DocxTools().set_format(p, bold=True, color="00FF00", size_pt=14)
    for r in p.runs:
        assert r.bold is True
        assert r.font.color.rgb == RGBColor(0x00, 0xFF, 0x00)
        assert r.font.size == Pt(14)
