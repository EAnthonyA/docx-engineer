"""
docxkit — run-aware helper toolkit for generated document-edit scripts.

All methods are single-pass over the document. The key primitive is
run-boundary-aware text replacement: python-docx may split a single visual
word across multiple Run objects, so naive `run.text.replace(...)` silently
misses matches that span runs. Every method here handles that correctly.
"""

import re
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _para_from_element(doc, el):
    """Wrap a raw lxml paragraph element in a python-docx Paragraph."""
    from docx.text.paragraph import Paragraph
    return Paragraph(el, doc)


def _replace_in_paragraph(para, find_fn):
    """
    Run the given find_fn on the concatenated text of para's runs.
    find_fn(full_text) must return a list of (start, end, replacement) tuples,
    where [start:end] in full_text should become replacement.

    Applies replacements right-to-left (to keep indices valid) by redistributing
    characters back into the original run sequence — preserving each run's
    formatting.

    Returns count of replacements made.
    """
    runs = para.runs
    if not runs:
        return 0

    # Build run map: character index → (run_index, char_index_within_run)
    run_starts = []
    full_text = []
    for ri, run in enumerate(runs):
        run_starts.append(len(full_text))
        full_text.extend(run.text)
    full_text = "".join(full_text)
    total = len(full_text)

    matches = find_fn(full_text)
    if not matches:
        return 0

    # Apply right-to-left so earlier indices stay valid
    matches = sorted(matches, key=lambda m: m[0], reverse=True)
    new_chars = list(full_text)
    for start, end, replacement in matches:
        new_chars[start:end] = list(replacement)

    # Re-distribute characters back into runs
    new_text = "".join(new_chars)

    if len(runs) == 1:
        runs[0].text = new_text
        return len(matches)

    # Compute new lengths proportionally based on original run lengths.
    # Strategy: keep each run's original char count; the last run absorbs remainder.
    cursor = 0
    for ri, run in enumerate(runs):
        orig_len = len(run.text)  # current (may be stale now) — use run_starts diff
        if ri + 1 < len(run_starts):
            orig_len = run_starts[ri + 1] - run_starts[ri]
        else:
            orig_len = total - run_starts[ri]

        if ri < len(runs) - 1:
            chunk = new_text[cursor:cursor + orig_len]
            run.text = chunk
            cursor += orig_len
        else:
            run.text = new_text[cursor:]

    return len(matches)


def _location_filter(locations):
    """Return (do_body, do_headers) booleans from locations string."""
    loc = (locations or "all").lower()
    do_body = loc in ("all", "body", "tables")
    do_hf = loc in ("all", "headers_footers")
    return do_body, do_hf


def _parse_color(color):
    """Hex string 'RRGGBB' or '#RRGGBB' -> RGBColor, or None."""
    if not color:
        return None
    c = color.lstrip("#")
    return RGBColor(int(c[0:2], 16), int(c[2:4], 16), int(c[4:6], 16))


# Per-character formatting is tracked as a 5-tuple:
#   (bold, italic, underline, color_rgb, size)
_NO_FMT = (None,) * 5


def _fmt_at(char_fmts, idx):
    """Format snapshot for character `idx`, or a neutral tuple if out of range."""
    return char_fmts[idx] if idx < len(char_fmts) else _NO_FMT


def _apply_fmt(run, fmt):
    """Apply a (bold, italic, underline, color, size) snapshot to a run."""
    bold, italic, underline, color, size = fmt
    run.bold = bold
    run.italic = italic
    run.underline = underline
    if color is not None:
        try:
            run.font.color.rgb = color
        except Exception:
            pass
    run.font.size = size


def _snapshot_char_fmts(runs):
    """Return (full_text, per-character format list) for a run sequence."""
    full_text = []
    char_fmts = []
    for run in runs:
        t = run.text
        full_text.append(t)
        try:
            c = run.font.color.rgb
        except Exception:
            c = None
        char_fmts.extend([(run.bold, run.italic, run.underline, c, run.font.size)] * len(t))
    return "".join(full_text), char_fmts


def _build_segments(full_text, matches):
    """Split full_text into (text, orig_pos, tagged) segments.

    tagged segments are the regex group(1) captures (to be reformatted); the
    rest is untouched text. Segments cover full_text left to right.
    """
    segments = []
    last_end = 0
    for m in matches:
        if m.start() > last_end:
            segments.append((full_text[last_end:m.start()], last_end, False))
        segments.append((m.group(1), m.start(1), True))
        last_end = m.end()
    if last_end < len(full_text):
        segments.append((full_text[last_end:], last_end, False))
    return segments


def _clear_runs(para):
    """Remove every child of the paragraph except its pPr.

    Handles runs wrapped in w:hyperlink / w:ins / w:del, where removing the
    run element directly would silently fail.
    """
    p_el = para._element
    for child in list(p_el):
        if child.tag != qn("w:pPr"):
            p_el.remove(child)


def _emit_preserving(para, text, char_fmts, orig_pos):
    """Add runs for `text`, grouping consecutive same-format chars into one run."""
    i = 0
    n = len(text)
    while i < n:
        fmt = _fmt_at(char_fmts, orig_pos + i)
        j = i + 1
        while j < n and _fmt_at(char_fmts, orig_pos + j) == fmt:
            j += 1
        _apply_fmt(para.add_run(text[i:j]), fmt)
        i = j


# ---------------------------------------------------------------------------
# Public toolkit class
# ---------------------------------------------------------------------------

class DocxTools:
    """
    Helper object passed as `tools` to the generated edit(doc, tools) function.

    All methods operate in-place on the document and return a count of changes.
    The `doc` parameter is always the python-docx Document instance.
    """

    # ------------------------------------------------------------------
    # Text replacement
    # ------------------------------------------------------------------

    def replace_text(self, doc, old: str, new: str, *, locations: str = "all") -> int:
        """
        Replace every literal occurrence of *old* with *new* across all paragraphs,
        handling matches that span run boundaries.

        locations: "all" | "body" | "tables" | "headers_footers"
        Returns total replacement count.
        """
        if not old:
            return 0

        def find(text):
            results = []
            start = 0
            while True:
                idx = text.find(old, start)
                if idx == -1:
                    break
                results.append((idx, idx + len(old), new))
                start = idx + len(old)
            return results

        return self._apply_to_paragraphs(doc, find, locations)

    def format_tagged(self, doc, pattern: str, *, flags: int = 0, locations: str = "all",
                      bold=None, italic=None, underline=None,
                      color: str = None, size_pt=None) -> int:
        """
        Find regex matches of pattern (must have one capture group), strip the full
        match, keep group(1), and apply the given formatting to the kept content.
        Run-boundary-aware; preserves original formatting on untouched text.

        Use this whenever you need to remove markup/tags AND format the content inside.
        Returns total count of replacements.
        """
        compiled = re.compile(pattern, flags)
        rgb = _parse_color(color)
        pt = Pt(size_pt) if size_pt is not None else None
        count = 0

        def _process(para):
            nonlocal count
            runs = para.runs
            if not runs:
                return

            full_text, char_fmts = _snapshot_char_fmts(runs)
            matches = list(compiled.finditer(full_text))
            if not matches:
                return

            _clear_runs(para)

            for text, orig_pos, tagged in _build_segments(full_text, matches):
                if not text:
                    continue
                if tagged:
                    # Tagged region: merge caller-supplied formatting over the original.
                    ob, oi, ou, oc, osz = _fmt_at(char_fmts, orig_pos)
                    _apply_fmt(para.add_run(text), (
                        bold if bold is not None else ob,
                        italic if italic is not None else oi,
                        underline if underline is not None else ou,
                        rgb if rgb is not None else oc,
                        pt if pt is not None else osz,
                    ))
                else:
                    _emit_preserving(para, text, char_fmts, orig_pos)

            count += len(matches)

        for para in self.iter_paragraphs(doc, locations=locations):
            _process(para)
        return count

    def regex_replace(self, doc, pattern: str, repl, *, flags: int = 0, locations: str = "all") -> int:
        """
        Replace every regex match of *pattern* with *repl* across all paragraphs,
        handling run-boundary spans.

        repl: str or callable(match) -> str
        locations: "all" | "body" | "tables" | "headers_footers"
        Returns total replacement count.
        """
        compiled = re.compile(pattern, flags)

        def find(text):
            results = []
            for m in compiled.finditer(text):
                replacement = repl(m) if callable(repl) else compiled.sub(repl, m.group())
                results.append((m.start(), m.end(), replacement))
            return results

        return self._apply_to_paragraphs(doc, find, locations)

    def _apply_to_paragraphs(self, doc, find_fn, locations):
        return sum(_replace_in_paragraph(para, find_fn)
                   for para in self.iter_paragraphs(doc, locations=locations))

    # ------------------------------------------------------------------
    # Paragraph iteration
    # ------------------------------------------------------------------

    def iter_paragraphs(self, doc, *, locations: str = "all"):
        """
        Generator yielding every paragraph in the document.
        Covers body paragraphs, table cell paragraphs, and header/footer paragraphs.

        locations: "all" | "body" | "tables" | "headers_footers"
        """
        do_body, do_hf = _location_filter(locations)

        if do_body:
            for block in doc.element.body:
                tag = block.tag.split("}")[-1]
                if tag == "p":
                    yield _para_from_element(doc, block)
                elif tag == "tbl":
                    for p_el in block.iter(qn("w:p")):
                        yield _para_from_element(doc, p_el)

        if do_hf:
            for section in doc.sections:
                for hf in (
                    section.header,
                    section.footer,
                    section.even_page_header,
                    section.even_page_footer,
                    section.first_page_header,
                    section.first_page_footer,
                ):
                    if hf is None:
                        continue
                    try:
                        for para in hf.paragraphs:
                            yield para
                    except Exception:
                        pass

    # ------------------------------------------------------------------
    # Paragraph deletion
    # ------------------------------------------------------------------

    def delete_paragraphs(self, doc, predicate) -> int:
        """
        Remove all paragraphs where predicate(paragraph.text) is truthy.
        Works in body and table cells (headers/footers excluded for safety).

        Returns count of removed paragraphs.
        """
        to_remove = []
        for para in self.iter_paragraphs(doc, locations="body"):
            try:
                if predicate(para.text):
                    to_remove.append(para._element)
            except Exception:
                pass

        for el in to_remove:
            parent = el.getparent()
            if parent is not None:
                parent.remove(el)

        return len(to_remove)

    # ------------------------------------------------------------------
    # Formatting
    # ------------------------------------------------------------------

    def set_format(self, target, *, bold=None, italic=None, underline=None,
                   color: str = None, size_pt=None) -> None:
        """
        Apply run-level formatting to a paragraph or a single run.

        target: python-docx Paragraph or Run
        color: hex string like "FF0000" (no #) or "RRGGBB"
        size_pt: font size in points (int or float)

        When target is a Paragraph, applies to every run in the paragraph.
        """
        runs = target.runs if hasattr(target, "runs") else [target]

        rgb = _parse_color(color)
        pt = Pt(size_pt) if size_pt is not None else None

        for run in runs:
            if bold is not None:
                run.bold = bold
            if italic is not None:
                run.italic = italic
            if underline is not None:
                run.underline = underline
            if rgb is not None:
                run.font.color.rgb = rgb
            if pt is not None:
                run.font.size = pt

    def set_style(self, paragraph, name: str) -> None:
        """
        Set the paragraph style by name. No-op if the style doesn't exist.
        """
        try:
            paragraph.style = name
        except Exception:
            pass
